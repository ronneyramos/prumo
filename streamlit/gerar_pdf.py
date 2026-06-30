"""
Geração de PDFs para o ERP MBR Engenharia.
  gerar_bm(dados, itens=None)            -> bytes  (Boletim de Medição)
  gerar_folha_pagamento(dados)           -> bytes  (Folha de Pagamento)
"""
from io import BytesIO
from datetime import date

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── Paleta MBR ────────────────────────────────────────────────────────────────
AZUL       = colors.HexColor("#2B59C3")
AZUL_ESC   = colors.HexColor("#1A3A8A")
CINZA_CAB  = colors.HexColor("#F0F2F6")
CINZA_LIN  = colors.HexColor("#F7F8FA")
BORDA      = colors.HexColor("#D0D5DD")
BRANCO     = colors.white
PRETO      = colors.black


def _fmt_brl(v) -> str:
    """Formata valor em R$ no padrão brasileiro."""
    try:
        f = f"{float(v):,.2f}"
        return "R$ " + f.replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "—"


def _pct(v) -> str:
    try:
        return f"{float(v):.2f}%"
    except Exception:
        return "—"


def _num(v, dec=2) -> str:
    try:
        return f"{float(v):,.{dec}f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "—"


# ── Estilos de parágrafo reutilizáveis ───────────────────────────────────────
def _estilos():
    s = getSampleStyleSheet()
    base = dict(fontName="Helvetica", fontSize=8, leading=10)
    return {
        "titulo":   ParagraphStyle("titulo",   fontName="Helvetica-Bold", fontSize=14,
                                   textColor=AZUL, alignment=TA_CENTER, leading=18),
        "subtitulo":ParagraphStyle("subtitulo",fontName="Helvetica-Bold", fontSize=10,
                                   textColor=AZUL_ESC, alignment=TA_CENTER, leading=14),
        "empresa":  ParagraphStyle("empresa",  fontName="Helvetica-Bold", fontSize=11,
                                   textColor=AZUL, alignment=TA_LEFT, leading=14),
        "cab_val":  ParagraphStyle("cab_val",  fontName="Helvetica-Bold", fontSize=8,
                                   textColor=PRETO, alignment=TA_LEFT, leading=10),
        "cab_label":ParagraphStyle("cab_label",fontName="Helvetica",      fontSize=7,
                                   textColor=colors.HexColor("#666666"), alignment=TA_LEFT, leading=9),
        "th":       ParagraphStyle("th",       fontName="Helvetica-Bold", fontSize=6.5,
                                   textColor=BRANCO, alignment=TA_CENTER, leading=8),
        "th_sub":   ParagraphStyle("th_sub",   fontName="Helvetica-Bold", fontSize=6,
                                   textColor=PRETO, alignment=TA_CENTER, leading=7.5),
        "td":       ParagraphStyle("td",       fontName="Helvetica",      fontSize=6.5,
                                   textColor=PRETO, alignment=TA_LEFT, leading=8),
        "td_r":     ParagraphStyle("td_r",     fontName="Helvetica",      fontSize=6.5,
                                   textColor=PRETO, alignment=TA_RIGHT, leading=8),
        "td_c":     ParagraphStyle("td_c",     fontName="Helvetica",      fontSize=6.5,
                                   textColor=PRETO, alignment=TA_CENTER, leading=8),
        "rodape":   ParagraphStyle("rodape",   fontName="Helvetica",      fontSize=7,
                                   textColor=colors.HexColor("#888888"), alignment=TA_CENTER),
        "total_l":  ParagraphStyle("total_l",  fontName="Helvetica-Bold", fontSize=8,
                                   textColor=PRETO, alignment=TA_LEFT, leading=10),
        "total_v":  ParagraphStyle("total_v",  fontName="Helvetica-Bold", fontSize=8,
                                   textColor=AZUL, alignment=TA_RIGHT, leading=10),
        "assin":    ParagraphStyle("assin",    fontName="Helvetica",      fontSize=8,
                                   textColor=PRETO, alignment=TA_CENTER, leading=10),
        "assin_n":  ParagraphStyle("assin_n",  fontName="Helvetica-Bold", fontSize=8,
                                   textColor=PRETO, alignment=TA_CENTER, leading=10),
    }


# ════════════════════════════════════════════════════════════════════════════════
# BOLETIM DE MEDIÇÃO
# ════════════════════════════════════════════════════════════════════════════════

def gerar_bm(dados: dict, itens: list = None) -> bytes:
    """
    Gera PDF do Boletim de Medição no layout real da MBR Engenharia.

    dados (obrigatórios):
        obra, cliente, periodo, num_bm,
        pct_anterior, pct_periodo, pct_acumulado,
        valor_contrato, valor_anterior, valor_periodo, valor_acumulado

    dados (opcionais):
        contratada  (default "MBR ENGENHARIA LTDA")

    itens (opcional): lista de dicts do orçamento com chaves
        codigo, descricao, unidade, quantidade, preco_venda, total_venda
        + pct_acumulado_item (opcional; usa pct_acumulado da obra como fallback)
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1.2*cm, rightMargin=1.2*cm,
        topMargin=1.5*cm,  bottomMargin=1.5*cm,
    )
    st = _estilos()
    hoje = date.today().strftime("%d/%m/%Y")
    w = landscape(A4)[0] - 2.4*cm   # largura útil

    # ── Extrair campos ────────────────────────────────────────────────────────
    obra        = dados.get("obra",          "—")
    cliente     = dados.get("cliente",       "—")
    contratada  = dados.get("contratada",    "MBR ENGENHARIA LTDA")
    periodo     = dados.get("periodo",       "—")
    num_bm      = dados.get("num_bm",        "—")
    pct_ant     = float(dados.get("pct_anterior",  0) or 0)
    pct_per     = float(dados.get("pct_periodo",   0) or 0)
    pct_acum    = float(dados.get("pct_acumulado", 0) or 0)
    val_contrato= float(dados.get("valor_contrato", 0) or 0)
    val_ant     = float(dados.get("valor_anterior", 0) or 0)
    val_per     = float(dados.get("valor_periodo",  0) or 0)
    val_acum    = float(dados.get("valor_acumulado",0) or 0)
    val_saldo   = val_contrato - val_acum

    story = []

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    def _cabecalho():
        # Linha 1: empresa | título | nº BM
        cab_top = Table(
            [[
                Paragraph(contratada, st["empresa"]),
                Paragraph(f"BOLETIM DE MEDIÇÃO — BM {num_bm:02d}" if str(num_bm).isdigit()
                          else f"BOLETIM DE MEDIÇÃO — BM {num_bm}", st["titulo"]),
                Paragraph(f"Emissão: {hoje}", st["cab_label"]),
            ]],
            colWidths=[w*0.28, w*0.44, w*0.28],
        )
        cab_top.setStyle(TableStyle([
            ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
            ("ALIGN",      (2,0), (2,0),   "RIGHT"),
            ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ]))
        story.append(cab_top)

        # Linha divisória
        story.append(HRFlowable(width="100%", thickness=2, color=AZUL, spaceAfter=4))

        # Linha 2: campos de identificação
        campos = [
            ("OBRA",        obra),
            ("CLIENTE",     cliente),
            ("CONTRATADA",  contratada),
            ("PERÍODO",     periodo),
            ("Nº BOLETIM",  f"BM-{num_bm:02d}" if str(num_bm).isdigit() else f"BM-{num_bm}"),
        ]
        col_w = w / len(campos)
        row_labels = [Paragraph(k, st["cab_label"]) for k,_ in campos]
        row_vals   = [Paragraph(v, st["cab_val"])   for _,v in campos]
        cab_info = Table([row_labels, row_vals], colWidths=[col_w]*len(campos))
        cab_info.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  CINZA_CAB),
            ("GRID",        (0,0), (-1,-1), 0.3, BORDA),
            ("TOPPADDING",  (0,0), (-1,-1), 3),
            ("BOTTOMPADDING",(0,0),(-1,-1), 3),
            ("LEFTPADDING", (0,0), (-1,-1), 4),
        ]))
        story.append(cab_info)
        story.append(Spacer(1, 4*mm))

    _cabecalho()

    # ── Tabela detalhada com itens do orçamento ────────────────────────────────
    if itens:
        _tabela_detalhada(story, st, w, itens, pct_per, pct_ant, pct_acum,
                          val_per, val_ant, val_acum, val_contrato)
    else:
        _tabela_resumo(story, st, w, pct_ant, pct_per, pct_acum,
                       val_ant, val_per, val_acum, val_contrato)

    # ── Total e assinatura ────────────────────────────────────────────────────
    story.append(Spacer(1, 3*mm))

    total_box = Table(
        [[
            Paragraph("VALOR TOTAL DO BOLETIM DE MEDIÇÃO:", st["total_l"]),
            Paragraph(_fmt_brl(val_per), st["total_v"]),
        ]],
        colWidths=[w*0.75, w*0.25],
    )
    total_box.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,-1), CINZA_CAB),
        ("BOX",          (0,0), (-1,-1), 1.5, AZUL),
        ("TOPPADDING",   (0,0), (-1,-1), 6),
        ("BOTTOMPADDING",(0,0), (-1,-1), 6),
        ("LEFTPADDING",  (0,0), (-1,-1), 8),
        ("RIGHTPADDING", (-1,0),(-1,-1), 8),
    ]))
    story.append(total_box)

    story.append(Spacer(1, 10*mm))

    # Assinaturas
    assin_data = [[
        Paragraph("____________________________\nContratante\n" + cliente,   st["assin"]),
        Paragraph("____________________________\nContratada\n" + contratada, st["assin"]),
        Paragraph("____________________________\nFiscal de Contrato",        st["assin"]),
        Paragraph("____________________________\nResponsável Técnico",       st["assin"]),
    ]]
    assin_t = Table(assin_data, colWidths=[w/4]*4)
    assin_t.setStyle(TableStyle([
        ("ALIGN",       (0,0), (-1,-1), "CENTER"),
        ("VALIGN",      (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",  (0,0), (-1,-1), 8),
    ]))
    story.append(assin_t)

    # Rodapé
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph(
        f"MBR Engenharia Ltda — Boletim de Medição BM-{num_bm} — Período: {periodo} — "
        f"Gerado em {hoje} pelo ERP MBR",
        st["rodape"]
    ))

    doc.build(story)
    return buf.getvalue()


def _tabela_detalhada(story, st, w, itens, pct_per, pct_ant, pct_acum,
                      val_per, val_ant, val_acum, val_contrato):
    """Tabela completa com itens do orçamento e colunas de executado."""
    # Larguras das colunas (total = w)
    # ITEM | DESCRIÇÃO | UND | QTDE | PR UNIT | PR TOT | QT.ANT | QT.PER | QT.ACUM | QT.SALDO | VL.ANT | VL.PER | VL.ACUM | VL.SALDO | %EXEC
    cws = [
        w*0.040,   # ITEM
        w*0.200,   # DESCRIÇÃO
        w*0.032,   # UND
        w*0.042,   # QTDE
        w*0.062,   # PR UNIT
        w*0.062,   # PR TOTAL
        w*0.046,   # QT ANT
        w*0.046,   # QT PER
        w*0.046,   # QT ACUM
        w*0.046,   # QT SALDO
        w*0.060,   # VL ANT
        w*0.060,   # VL PER
        w*0.060,   # VL ACUM
        w*0.060,   # VL SALDO
        w*0.038,   # % EXEC
    ]

    P = Paragraph

    # Row 0: cabeçalho de grupo
    row0 = [
        P("DISCRIMINAÇÃO DOS SERVIÇOS", st["th"]),   # 0 (span 0-2)
        "", "",
        P("PREVISTO NO ORÇAMENTO",      st["th"]),   # 3 (span 3-5)
        "", "",
        P("EXECUTADO FÍSICO (QTDE)",    st["th"]),   # 6 (span 6-9)
        "", "", "",
        P("EXECUTADO FINANCEIRO (R$)",  st["th"]),   # 10 (span 10-13)
        "", "", "",
        P("DESVIO",                     st["th"]),   # 14
    ]

    # Row 1: sub-cabeçalhos
    row1 = [
        P("ITEM",        st["th_sub"]),
        P("REFERÊNCIA / SERVIÇO", st["th_sub"]),
        P("UND",         st["th_sub"]),
        P("QTDES",       st["th_sub"]),
        P("PR UNIT\nINCL. BDI",   st["th_sub"]),
        P("PR TOTAL\nINCL. BDI",  st["th_sub"]),
        P("QUANT.\nANTERIOR",     st["th_sub"]),
        P("QUANT.\nPERÍODO",      st["th_sub"]),
        P("QUANT.\nACUMULADA",    st["th_sub"]),
        P("QUANT.\nSALDO",        st["th_sub"]),
        P("ACUM. ATÉ\nPER. ANT.", st["th_sub"]),
        P("VALOR\nPERÍODO",       st["th_sub"]),
        P("VALOR\nACUMULADO",     st["th_sub"]),
        P("PREÇO\nSALDO",         st["th_sub"]),
        P("(%)\nEXEC.",            st["th_sub"]),
    ]

    data = [row0, row1]

    total_previsto = 0.0

    for r in itens:
        cod     = str(r.get("codigo","") or "")
        desc    = str(r.get("descricao","") or "")
        und     = str(r.get("unidade","") or "")
        qtde    = float(r.get("quantidade")    or 0)
        pr_unit = float(r.get("preco_venda")   or 0)
        pr_tot  = float(r.get("total_venda")   or 0)
        total_previsto += pr_tot

        # Executado físico (proporcional ao % do período)
        qt_ant   = qtde * pct_ant  / 100
        qt_per   = qtde * pct_per  / 100
        qt_acum  = qtde * pct_acum / 100
        qt_saldo = max(0.0, qtde - qt_acum)

        # Executado financeiro
        vl_ant   = pr_tot * pct_ant  / 100
        vl_per   = pr_tot * pct_per  / 100
        vl_acum  = pr_tot * pct_acum / 100
        vl_saldo = max(0.0, pr_tot - vl_acum)
        pct_exec = pct_acum

        data.append([
            P(cod,              st["td_c"]),
            P(desc,             st["td"]),
            P(und,              st["td_c"]),
            P(_num(qtde),       st["td_r"]),
            P(_num(pr_unit),    st["td_r"]),
            P(_num(pr_tot),     st["td_r"]),
            P(_num(qt_ant),     st["td_r"]),
            P(_num(qt_per),     st["td_r"]),
            P(_num(qt_acum),    st["td_r"]),
            P(_num(qt_saldo),   st["td_r"]),
            P(_num(vl_ant),     st["td_r"]),
            P(_num(vl_per),     st["td_r"]),
            P(_num(vl_acum),    st["td_r"]),
            P(_num(vl_saldo),   st["td_r"]),
            P(_pct(pct_exec),   st["td_c"]),
        ])

    # Linha de totais
    tot_ant  = total_previsto * pct_ant  / 100
    tot_per  = total_previsto * pct_per  / 100
    tot_acum = total_previsto * pct_acum / 100
    tot_sal  = max(0.0, total_previsto - tot_acum)
    data.append([
        P("TOTAL", st["th_sub"]), "",
        "", P("—", st["td_c"]),
        P("—", st["td_c"]),
        P(_num(total_previsto), st["td_r"]),
        P("—", st["td_c"]),
        P("—", st["td_c"]),
        P("—", st["td_c"]),
        P("—", st["td_c"]),
        P(_num(tot_ant),  st["td_r"]),
        P(_num(tot_per),  st["td_r"]),
        P(_num(tot_acum), st["td_r"]),
        P(_num(tot_sal),  st["td_r"]),
        P(_pct(pct_acum), st["td_c"]),
    ])

    tbl = Table(data, colWidths=cws, repeatRows=2)

    n_data  = len(data)
    n_itens = n_data - 3   # exclui 2 headers + 1 total

    style = [
        # ── Spans de cabeçalho de grupo (row 0) ──────────────────────────────
        ("SPAN",        (0, 0), (2, 0)),   # DISCRIMINAÇÃO
        ("SPAN",        (3, 0), (5, 0)),   # PREVISTO
        ("SPAN",        (6, 0), (9, 0)),   # EXEC FÍSICO
        ("SPAN",        (10,0), (13,0)),   # EXEC FINANCEIRO
        # col 14 sem span (DESVIO)

        # ── Span de total ──────────────────────────────────────────────────────
        ("SPAN",        (0, n_data-1), (2, n_data-1)),

        # ── Cores de fundo dos cabeçalhos de grupo ────────────────────────────
        ("BACKGROUND",  (0, 0), (2, 0),   AZUL_ESC),
        ("BACKGROUND",  (3, 0), (5, 0),   AZUL),
        ("BACKGROUND",  (6, 0), (9, 0),   colors.HexColor("#1E7D4A")),
        ("BACKGROUND",  (10,0), (13,0),   colors.HexColor("#7B3F9E")),
        ("BACKGROUND",  (14,0), (14,0),   colors.HexColor("#B84A00")),

        # ── Sub-cabeçalhos (row 1) ─────────────────────────────────────────────
        ("BACKGROUND",  (0, 1), (-1, 1),  CINZA_CAB),
        ("TEXTCOLOR",   (0, 1), (-1, 1),  colors.HexColor("#222222")),

        # ── Linha de totais ───────────────────────────────────────────────────
        ("BACKGROUND",  (0, n_data-1), (-1, n_data-1), CINZA_CAB),
        ("FONTNAME",    (0, n_data-1), (-1, n_data-1), "Helvetica-Bold"),
        ("FONTSIZE",    (0, n_data-1), (-1, n_data-1), 6.5),

        # ── Grade e alinhamento gerais ────────────────────────────────────────
        ("GRID",        (0, 0), (-1, -1), 0.3, BORDA),
        ("ALIGN",       (0, 0), (-1, 0),  "CENTER"),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",  (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0,0), (-1,-1),  2),
        ("LEFTPADDING", (0, 0), (-1, -1), 2),
        ("RIGHTPADDING",(0, 0), (-1, -1), 2),

        # ── Linhas alternadas nos dados ───────────────────────────────────────
    ]

    # Linhas alternadas
    for i in range(2, n_data - 1):
        if i % 2 == 0:
            style.append(("BACKGROUND", (0, i), (-1, i), CINZA_LIN))

    tbl.setStyle(TableStyle(style))
    story.append(tbl)


def _tabela_resumo(story, st, w, pct_ant, pct_per, pct_acum,
                   val_ant, val_per, val_acum, val_contrato):
    """Tabela de resumo financeiro quando não há itens de orçamento carregados."""
    story.append(Paragraph("Resumo Financeiro do Boletim", _estilos()["subtitulo"]))
    story.append(Spacer(1, 3*mm))

    val_saldo = max(0.0, val_contrato - val_acum)

    rows = [
        ["DESCRIÇÃO",                        "% FÍSICO",      "VALOR (R$)"],
        ["Valor do Contrato",                 "100,00%",        _fmt_brl(val_contrato)],
        ["Executado até Período Anterior",    _pct(pct_ant),    _fmt_brl(val_ant)],
        ["Executado no Período",              _pct(pct_per),    _fmt_brl(val_per)],
        ["Executado Acumulado",               _pct(pct_acum),   _fmt_brl(val_acum)],
        ["Saldo a Executar",                  _pct(100 - pct_acum), _fmt_brl(val_saldo)],
    ]

    col_w = [w*0.55, w*0.20, w*0.25]
    tbl = Table(rows, colWidths=col_w)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0),   AZUL),
        ("TEXTCOLOR",    (0,0), (-1,0),   BRANCO),
        ("FONTNAME",     (0,0), (-1,0),   "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1),  9),
        ("ALIGN",        (0,0), (0,-1),   "LEFT"),
        ("ALIGN",        (1,0), (-1,-1),  "CENTER"),
        ("BACKGROUND",   (0,3), (-1,3),   colors.HexColor("#EAF0FF")),
        ("BACKGROUND",   (0,4), (-1,4),   colors.HexColor("#E6F4EA")),
        ("FONTNAME",     (0,4), (-1,4),   "Helvetica-Bold"),
        ("FONTNAME",     (0,3), (-1,3),   "Helvetica-Bold"),
        ("GRID",         (0,0), (-1,-1),  0.5, BORDA),
        ("TOPPADDING",   (0,0), (-1,-1),  5),
        ("BOTTOMPADDING",(0,0), (-1,-1),  5),
        ("LEFTPADDING",  (0,0), (-1,-1),  8),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),  [BRANCO, CINZA_LIN]),
    ]))
    story.append(tbl)


# ════════════════════════════════════════════════════════════════════════════════
# FOLHA DE PAGAMENTO
# ════════════════════════════════════════════════════════════════════════════════

def gerar_folha_pagamento(dados: dict) -> bytes:
    """
    Gera PDF da Folha de Pagamento.

    dados:
        ref_mes      str        (ex: "06/2026")
        obra         str
        colaboradores list[dict] com chaves: nome, cargo, tipo_contrato,
                                             salario, inss, fgts, liquido
        bruto        float
        inss_patronal float
        rat          float
        fgts         float
        custo_empresa float
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=2.0*cm,  bottomMargin=2.0*cm,
    )
    st = _estilos()
    hoje = date.today().strftime("%d/%m/%Y")
    w = A4[0] - 3.0*cm

    ref_mes       = dados.get("ref_mes",       "—")
    obra          = dados.get("obra",          "Geral")
    colaboradores = dados.get("colaboradores", [])
    bruto         = float(dados.get("bruto",          0) or 0)
    inss_pat      = float(dados.get("inss_patronal",  0) or 0)
    rat           = float(dados.get("rat",            0) or 0)
    fgts_emp      = float(dados.get("fgts",           0) or 0)
    custo_emp     = float(dados.get("custo_empresa",  0) or 0)

    story = []

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    story.append(Paragraph("MBR ENGENHARIA LTDA", st["empresa"]))
    story.append(Spacer(1, 1*mm))
    story.append(HRFlowable(width="100%", thickness=2, color=AZUL, spaceAfter=4))

    story.append(Paragraph("FOLHA DE PAGAMENTO", st["titulo"]))
    story.append(Spacer(1, 2*mm))

    # Campos de identificação
    campos = [
        ("MÊS DE REFERÊNCIA", ref_mes),
        ("OBRA",               obra),
        ("TOTAL DE COLABORADORES", str(len(colaboradores))),
        ("DATA DE EMISSÃO",    hoje),
    ]
    col_w_cab = w / len(campos)
    cab_tbl = Table(
        [[Paragraph(k, st["cab_label"]) for k,_ in campos],
         [Paragraph(v, st["cab_val"])   for _,v in campos]],
        colWidths=[col_w_cab]*len(campos),
    )
    cab_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0),  CINZA_CAB),
        ("GRID",         (0,0), (-1,-1), 0.3, BORDA),
        ("TOPPADDING",   (0,0), (-1,-1), 3),
        ("BOTTOMPADDING",(0,0), (-1,-1), 3),
        ("LEFTPADDING",  (0,0), (-1,-1), 6),
    ]))
    story.append(cab_tbl)
    story.append(Spacer(1, 5*mm))

    # ── Tabela de colaboradores ────────────────────────────────────────────────
    story.append(Paragraph("Discriminação por Colaborador", st["subtitulo"]))
    story.append(Spacer(1, 2*mm))

    P = Paragraph
    header = [
        P("Nº",           st["th"]),
        P("NOME",         st["th"]),
        P("CARGO",        st["th"]),
        P("CONTRATO",     st["th"]),
        P("SALÁRIO\nBRUTO (R$)", st["th"]),
        P("INSS\nEMPREGADO (R$)", st["th"]),
        P("FGTS\nEMP. (R$)", st["th"]),
        P("LÍQUIDO (R$)", st["th"]),
    ]

    cws_f = [w*0.05, w*0.23, w*0.15, w*0.09,
             w*0.12, w*0.12,  w*0.10, w*0.14]

    rows_f = [header]
    for i, c in enumerate(colaboradores, 1):
        sal = float(c.get("salario", 0) or 0)
        inss = float(c.get("inss",   0) or 0)
        fgts = float(c.get("fgts",   0) or 0)
        liq  = float(c.get("liquido",0) or 0)
        rows_f.append([
            P(str(i),                 st["td_c"]),
            P(str(c.get("nome","—")), st["td"]),
            P(str(c.get("cargo","—")),st["td"]),
            P(str(c.get("tipo_contrato","CLT")), st["td_c"]),
            P(_fmt_brl(sal),          st["td_r"]),
            P(_fmt_brl(inss),         st["td_r"]),
            P(_fmt_brl(fgts),         st["td_r"]),
            P(_fmt_brl(liq),          st["td_r"]),
        ])

    n_rows = len(rows_f)
    tbl_f = Table(rows_f, colWidths=cws_f, repeatRows=1)
    style_f = [
        ("BACKGROUND",   (0,0), (-1,0),  AZUL),
        ("TEXTCOLOR",    (0,0), (-1,0),  BRANCO),
        ("GRID",         (0,0), (-1,-1), 0.3, BORDA),
        ("ALIGN",        (0,0), (-1,0),  "CENTER"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",   (0,0), (-1,-1), 3),
        ("BOTTOMPADDING",(0,0), (-1,-1), 3),
        ("LEFTPADDING",  (0,0), (-1,-1), 3),
        ("RIGHTPADDING", (0,0), (-1,-1), 3),
    ]
    for i in range(1, n_rows):
        if i % 2 == 0:
            style_f.append(("BACKGROUND", (0, i), (-1, i), CINZA_LIN))
    tbl_f.setStyle(TableStyle(style_f))
    story.append(tbl_f)

    story.append(Spacer(1, 5*mm))

    # ── Totais + Encargos da Empresa ─────────────────────────────────────────
    story.append(Paragraph("Encargos e Custo Total da Empresa", st["subtitulo"]))
    story.append(Spacer(1, 2*mm))

    enc_rows = [
        [P("DESCRIÇÃO",                  st["th"]),
         P("ALÍQUOTA", st["th"]),
         P("VALOR (R$)",                 st["th"])],
        [P("Salário Bruto Total",         st["td"]),  P("—",      st["td_c"]), P(_fmt_brl(bruto),    st["td_r"])],
        [P("INSS Patronal",               st["td"]),  P("20,0%",  st["td_c"]), P(_fmt_brl(inss_pat), st["td_r"])],
        [P("RAT (Risco Acidente de Trabalho)", st["td"]), P("3,0%", st["td_c"]), P(_fmt_brl(rat),   st["td_r"])],
        [P("FGTS (Empresa)",              st["td"]),  P("8,0%",   st["td_c"]), P(_fmt_brl(fgts_emp), st["td_r"])],
        [P("CUSTO TOTAL EMPRESA",         st["th"]),  P("—",      st["td_c"]), P(_fmt_brl(custo_emp),st["total_v"])],
    ]
    cws_enc = [w*0.55, w*0.15, w*0.30]
    tbl_enc = Table(enc_rows, colWidths=cws_enc)
    tbl_enc.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,0),     AZUL),
        ("TEXTCOLOR",   (0,0), (-1,0),     BRANCO),
        ("BACKGROUND",  (0,-1),(-1,-1),    CINZA_CAB),
        ("FONTNAME",    (0,-1),(-1,-1),    "Helvetica-Bold"),
        ("FONTSIZE",    (0,-1),(0,-1),     8),
        ("TEXTCOLOR",   (2,-1),(2,-1),     AZUL),
        ("FONTSIZE",    (2,-1),(2,-1),     9),
        ("BOX",         (0,-1),(-1,-1),    1.5, AZUL),
        ("GRID",        (0,0), (-1,-1),    0.3, BORDA),
        ("ALIGN",       (1,0), (-1,-1),    "CENTER"),
        ("ALIGN",       (0,0), (0,-1),     "LEFT"),
        ("VALIGN",      (0,0), (-1,-1),    "MIDDLE"),
        ("TOPPADDING",  (0,0), (-1,-1),    4),
        ("BOTTOMPADDING",(0,0),(-1,-1),    4),
        ("LEFTPADDING", (0,0), (-1,-1),    6),
    ]))
    story.append(tbl_enc)

    story.append(Spacer(1, 10*mm))

    # ── Assinaturas ───────────────────────────────────────────────────────────
    assin = Table(
        [[
            P("____________________________\nResponsável Financeiro", st["assin"]),
            P("____________________________\nDiretor / Sócio",        st["assin"]),
            P("____________________________\nRH / Departamento Pessoal", st["assin"]),
        ]],
        colWidths=[w/3]*3,
    )
    assin.setStyle(TableStyle([
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("TOPPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(assin)

    story.append(Spacer(1, 4*mm))
    story.append(Paragraph(
        f"MBR Engenharia Ltda — Folha de Pagamento {ref_mes} — Obra: {obra} — "
        f"Gerado em {hoje} pelo ERP MBR",
        st["rodape"]
    ))

    doc.build(story)
    return buf.getvalue()


# ── DIÁRIO DE OBRA (RDO) ─────────────────────────────────────────────────────

def gerar_rdo(dados: dict, fotos: list = None) -> bytes:
    """PDF do Diário de Obra com relatório fotográfico."""
    import io, urllib.request
    from reportlab.lib.pagesizes import A4 as _A4
    from reportlab.lib import colors as _colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image as RLImage, HRFlowable)
    from reportlab.lib.enums import TA_CENTER

    AZUL = _colors.HexColor("#2B59C3")
    CINZA = _colors.HexColor("#F0F2F6")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=_A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    story = []

    s_titulo = ParagraphStyle("rt", fontSize=16, fontName="Helvetica-Bold",
                               textColor=AZUL, alignment=TA_CENTER, spaceAfter=4)
    s_sub    = ParagraphStyle("rs", fontSize=11, fontName="Helvetica",
                               alignment=TA_CENTER, spaceAfter=12)
    s_val    = ParagraphStyle("rv", fontSize=10, fontName="Helvetica", spaceAfter=4)
    s_sec    = ParagraphStyle("rsc", fontSize=11, fontName="Helvetica-Bold",
                               textColor=AZUL, spaceBefore=10, spaceAfter=4)

    story.append(Paragraph("MBR ENGENHARIA", s_titulo))
    story.append(Paragraph("DIÁRIO DE OBRA — RDO", s_sub))
    story.append(HRFlowable(width="100%", thickness=2, color=AZUL))
    story.append(Spacer(1, 0.3*cm))

    info_data = [
        ["Obra:",         str(dados.get("Obra","")),        "Data:",         str(dados.get("Data",""))],
        ["Responsável:",  str(dados.get("Responsável","")), "Status:",       str(dados.get("Status Dia","Normal"))],
        ["Clima Manhã:",  str(dados.get("Clima Manhã","")), "Clima Tarde:",  str(dados.get("Clima Tarde",""))],
        ["Efetivo Total:",f"{dados.get('Efetivo Total',0)} pessoas", "", ""],
    ]
    t_info = Table(info_data, colWidths=[3*cm, 7*cm, 3*cm, 4*cm])
    t_info.setStyle(TableStyle([
        ("FONTNAME",  (0,0),(-1,-1), "Helvetica"),
        ("FONTNAME",  (0,0),(0,-1),  "Helvetica-Bold"),
        ("FONTNAME",  (2,0),(2,-1),  "Helvetica-Bold"),
        ("FONTSIZE",  (0,0),(-1,-1), 9),
        ("ROWBACKGROUNDS",(0,0),(-1,-1),[_colors.white, CINZA]),
        ("GRID",      (0,0),(-1,-1), 0.5, _colors.grey),
        ("TOPPADDING",(0,0),(-1,-1), 4),
        ("BOTTOMPADDING",(0,0),(-1,-1), 4),
    ]))
    story.append(t_info)
    story.append(Spacer(1, 0.4*cm))

    def _sec(titulo, texto):
        if not texto or not str(texto).strip() or str(texto).strip() in ("nan","None"):
            return
        story.append(Paragraph(titulo, s_sec))
        story.append(Paragraph(str(texto).replace("\n","<br/>"), s_val))

    _sec("Atividades Realizadas",         dados.get("Atividades",""))
    _sec("Ocorrências / Intercorrências", dados.get("Ocorrências",""))
    _sec("Equipamentos Utilizados",       dados.get("Equipamentos",""))
    _sec("Observações Gerais",            dados.get("Observações",""))

    fotos = fotos or []
    if fotos:
        story.append(HRFlowable(width="100%", thickness=1, color=AZUL))
        story.append(Paragraph("RELATÓRIO FOTOGRÁFICO", s_sec))
        foto_items = []
        for foto in fotos:
            url  = foto.get("url","") if isinstance(foto, dict) else str(foto)
            nome = foto.get("nome","Foto") if isinstance(foto, dict) else "Foto"
            try:
                with urllib.request.urlopen(url, timeout=10) as resp:
                    img_bytes = io.BytesIO(resp.read())
                foto_items.append([RLImage(img_bytes, width=8*cm, height=6*cm),
                                   Paragraph(nome, s_val)])
            except Exception:
                foto_items.append([Paragraph(f"[{nome}]", s_val), ""])
        for i in range(0, len(foto_items), 2):
            par = foto_items[i:i+2]
            if len(par) == 1:
                par.append(["", ""])
            t_f = Table([[par[0][0], par[1][0]], [par[0][1], par[1][1]]],
                        colWidths=[8.5*cm, 8.5*cm])
            t_f.setStyle(TableStyle([
                ("ALIGN",(0,0),(-1,-1),"CENTER"),
                ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                ("TOPPADDING",(0,0),(-1,-1),4),
                ("BOTTOMPADDING",(0,0),(-1,-1),4),
            ]))
            story.append(t_f)
            story.append(Spacer(1, 0.3*cm))

    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=_colors.grey))
    t_ass = Table([
        ["_"*35,                                            "_"*35],
        [dados.get("Responsável","Engenheiro Responsável"), "Fiscal / Encarregado"],
    ], colWidths=[8.5*cm, 8.5*cm])
    t_ass.setStyle(TableStyle([
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("FONTNAME",(0,1),(-1,1),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),8),
        ("TOPPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(t_ass)

    doc.build(story)
    return buf.getvalue()


def gerar_rdo_docx(dados: dict, fotos: list = None) -> bytes:
    """DOCX do Diário de Obra com relatório fotográfico. Requer python-docx."""
    import io, urllib.request
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    h = doc.add_heading("MBR ENGENHARIA — DIÁRIO DE OBRA (RDO)", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"

    def _c(row, col, text, bold=False):
        cell = tbl.cell(row, col)
        cell.text = str(text)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    _c(0,0,"Obra:", True);        _c(0,1, dados.get("Obra",""))
    _c(0,2,"Data:", True);        _c(0,3, dados.get("Data",""))
    _c(1,0,"Responsável:", True); _c(1,1, dados.get("Responsável",""))
    _c(1,2,"Status:", True);      _c(1,3, dados.get("Status Dia",""))
    _c(2,0,"Clima Manhã:", True); _c(2,1, dados.get("Clima Manhã",""))
    _c(2,2,"Clima Tarde:", True); _c(2,3, dados.get("Clima Tarde",""))
    _c(3,0,"Efetivo Total:", True); _c(3,1, f"{dados.get('Efetivo Total',0)} pessoas")
    _c(3,2,""); _c(3,3,"")
    doc.add_paragraph()

    def _sd(titulo, texto):
        if not texto or not str(texto).strip() or str(texto).strip() in ("nan","None"):
            return
        doc.add_heading(titulo, 2)
        doc.add_paragraph(str(texto))

    _sd("Atividades Realizadas",         dados.get("Atividades",""))
    _sd("Ocorrências / Intercorrências", dados.get("Ocorrências",""))
    _sd("Equipamentos Utilizados",       dados.get("Equipamentos",""))
    _sd("Observações Gerais",            dados.get("Observações",""))

    fotos = fotos or []
    if fotos:
        doc.add_heading("RELATÓRIO FOTOGRÁFICO", 2)
        for foto in fotos:
            url  = foto.get("url","") if isinstance(foto, dict) else str(foto)
            nome = foto.get("nome","Foto") if isinstance(foto, dict) else "Foto"
            try:
                with urllib.request.urlopen(url, timeout=10) as resp:
                    img_bytes = io.BytesIO(resp.read())
                doc.add_picture(img_bytes, width=Cm(12))
                p = doc.add_paragraph(nome)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"[Foto: {nome}]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def gerar_relatorio_gerencial(dados: dict) -> bytes:
    """Gera PDF do Relatório Gerencial Mensal MBR Engenharia."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as _colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from datetime import date as _date
    import io as _io
    import pandas as _pd

    AZUL   = _colors.HexColor("#2B59C3")
    CINZA  = _colors.HexColor("#F0F2F6")
    BRANCO = _colors.white

    def _num(v):
        try:
            return float(str(v).replace("R$","").replace(" ","").replace(".","").replace(",","."))
        except:
            return 0.0

    def _brl(v):
        try:
            n = _num(v)
            return f"R$ {n:,.2f}".replace(",","X").replace(".",",").replace("X",".")
        except:
            return "R$ 0,00"

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    def _S(name, **kw):
        return ParagraphStyle(name, **kw)

    sTit = _S("tit", fontSize=20, fontName="Helvetica-Bold", textColor=BRANCO, alignment=TA_CENTER)
    sSub = _S("sub", fontSize=11, fontName="Helvetica",      textColor=BRANCO, alignment=TA_CENTER)
    sH2  = _S("h2",  fontSize=12, fontName="Helvetica-Bold", textColor=AZUL, spaceBefore=14, spaceAfter=4)

    def _hdr_style():
        return TableStyle([
            ("BACKGROUND",(0,0),(-1,0), AZUL), ("TEXTCOLOR",(0,0),(-1,0), BRANCO),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[BRANCO, CINZA]),
            ("GRID",(0,0),(-1,-1),0.5,_colors.HexColor("#DDDDDD")),
            ("TOPPADDING",(0,0),(-1,-1),5), ("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),8),
        ])

    story = []
    hoje  = _date.today().strftime("%d/%m/%Y")
    mes   = dados.get("mes_ref", _date.today().strftime("%B/%Y"))

    obras_df = dados.get("obras",    _pd.DataFrame())
    meds_df  = dados.get("medicoes", _pd.DataFrame())
    cp_df    = dados.get("contas_pagar",   _pd.DataFrame())
    cr_df    = dados.get("contas_receber", _pd.DataFrame())
    nc_df    = dados.get("ncs",            _pd.DataFrame())
    func_df  = dados.get("funcionarios",   _pd.DataFrame())

    # ── CAPA ──────────────────────────────────────────────────────────────────
    capa = Table([
        [Paragraph("MBR ENGENHARIA", sTit)],
        [Paragraph("RELATÓRIO GERENCIAL MENSAL", sTit)],
        [Spacer(1, 0.3*cm)],
        [Paragraph(f"Período de Referência: {mes}", sSub)],
        [Paragraph(f"Emitido em: {hoje}", sSub)],
    ], colWidths=[17*cm])
    capa.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1), AZUL),
        ("TOPPADDING",(0,0),(-1,-1), 18),
        ("BOTTOMPADDING",(0,0),(-1,-1), 18),
    ]))
    story += [Spacer(1, 2*cm), capa, Spacer(1, 1*cm)]

    # ── KPIs EXECUTIVOS ───────────────────────────────────────────────────────
    story.append(Paragraph("1. VISÃO EXECUTIVA", sH2))
    story.append(HRFlowable(width="100%", thickness=2, color=AZUL))
    story.append(Spacer(1, 0.3*cm))

    n_ativas = 0
    tot_cont = 0.0
    if not obras_df.empty:
        if "Status" in obras_df.columns:
            n_ativas = int((obras_df["Status"] == "Em andamento").sum())
        if "Valor Contrato (R$)" in obras_df.columns:
            tot_cont = obras_df["Valor Contrato (R$)"].apply(_num).sum()

    tot_med = 0.0
    if not meds_df.empty and "Valor Medido (R$)" in meds_df.columns:
        tot_med = meds_df["Valor Medido (R$)"].apply(_num).sum()

    tot_cp = 0.0
    if not cp_df.empty and "Valor (R$)" in cp_df.columns:
        tot_cp = cp_df["Valor (R$)"].apply(_num).sum()

    tot_cr = 0.0
    if not cr_df.empty and "Valor (R$)" in cr_df.columns:
        tot_cr = cr_df["Valor (R$)"].apply(_num).sum()

    n_func = len(func_df) if not func_df.empty else 0

    n_nc_ab = 0
    if not nc_df.empty and "Status" in nc_df.columns:
        n_nc_ab = int(nc_df["Status"].str.lower().isin(["aberta","em análise","pendente"]).sum())

    kpi_data = [
        ["INDICADOR", "VALOR"],
        ["Obras em andamento",           str(n_ativas)],
        ["Total contratado (portfólio)", _brl(tot_cont)],
        ["Total medido acumulado",       _brl(tot_med)],
        ["Saldo a medir",                _brl(max(0, tot_cont - tot_med))],
        ["Contas a pagar (total)",       _brl(tot_cp)],
        ["Contas a receber (total)",     _brl(tot_cr)],
        ["Resultado (Receber − Pagar)",  _brl(tot_cr - tot_cp)],
        ["Colaboradores ativos",         str(n_func)],
        ["NCs em aberto",                str(n_nc_ab)],
    ]
    t_kpi = Table(kpi_data, colWidths=[10*cm, 7*cm])
    ts = _hdr_style()
    ts.add("FONTNAME",(0,1),(0,-1),"Helvetica-Bold")
    t_kpi.setStyle(ts)
    story += [t_kpi, Spacer(1, 0.5*cm)]

    # ── STATUS DAS OBRAS ──────────────────────────────────────────────────────
    if not obras_df.empty:
        story.append(Paragraph("2. STATUS DAS OBRAS", sH2))
        story.append(HRFlowable(width="100%", thickness=2, color=AZUL))
        story.append(Spacer(1, 0.3*cm))

        ob_data = [["OBRA","STATUS","AVANÇO","VALOR CONTRATO","RESPONSÁVEL"]]
        for _, o in obras_df.iterrows():
            pct = _num(o.get("% Físico", 0))
            sem = "🟢" if pct >= 70 else "🟡" if pct >= 40 else "🔴"
            ob_data.append([
                str(o.get("Nome",""))[:30],
                str(o.get("Status","")),
                f"{sem} {pct:.1f}%",
                _brl(o.get("Valor Contrato (R$)", 0)),
                str(o.get("Responsável",""))[:20],
            ])
        t_ob = Table(ob_data, colWidths=[5*cm, 3*cm, 2.5*cm, 3.5*cm, 3*cm])
        t_ob.setStyle(_hdr_style())
        story += [t_ob, Spacer(1, 0.4*cm)]

    # ── RESUMO FINANCEIRO ─────────────────────────────────────────────────────
    story.append(Paragraph("3. RESUMO FINANCEIRO", sH2))
    story.append(HRFlowable(width="100%", thickness=2, color=AZUL))
    story.append(Spacer(1, 0.3*cm))

    fin_data = [
        ["ITEM", "VALOR", "SITUAÇÃO"],
        ["Contas a Receber (total)", _brl(tot_cr), ""],
        ["Contas a Pagar (total)",   _brl(tot_cp), ""],
        ["Saldo (Receber − Pagar)",  _brl(tot_cr - tot_cp),
         "Positivo" if tot_cr >= tot_cp else "Negativo"],
    ]
    t_fin = Table(fin_data, colWidths=[7*cm, 5*cm, 5*cm])
    ts_fin = _hdr_style()
    ts_fin.add("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold")
    t_fin.setStyle(ts_fin)
    story += [t_fin, Spacer(1, 0.4*cm)]

    # ── NÃO-CONFORMIDADES ─────────────────────────────────────────────────────
    if not nc_df.empty:
        story.append(Paragraph("4. NÃO-CONFORMIDADES", sH2))
        story.append(HRFlowable(width="100%", thickness=2, color=AZUL))
        story.append(Spacer(1, 0.3*cm))

        nc_data = [["OBRA","DESCRIÇÃO","GRAVIDADE","STATUS"]]
        for _, r in nc_df.iterrows():
            nc_data.append([
                str(r.get("Obra",""))[:25],
                str(r.get("Descrição",""))[:40],
                str(r.get("Gravidade","")),
                str(r.get("Status","")),
            ])
        t_nc = Table(nc_data, colWidths=[4*cm, 7*cm, 3*cm, 3*cm])
        t_nc.setStyle(_hdr_style())
        story += [t_nc, Spacer(1, 0.4*cm)]

    # ── ASSINATURA ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=_colors.grey))
    t_ass = Table([
        ["_"*40, "_"*40],
        ["Responsável Técnico / Engenheiro", "Aprovação / Contratante"],
    ], colWidths=[8.5*cm, 8.5*cm])
    t_ass.setStyle(TableStyle([
        ("ALIGN",(0,0),(-1,-1),"CENTER"),
        ("FONTSIZE",(0,0),(-1,-1),8),
        ("TOPPADDING",(0,0),(-1,-1),8),
    ]))
    story.append(t_ass)

    doc.build(story)
    return buf.getvalue()
